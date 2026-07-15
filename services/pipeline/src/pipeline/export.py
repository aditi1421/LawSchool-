"""Export artifacts to a clean .docx brief."""

from io import BytesIO

from docx import Document
from docx.shared import Pt

from pipeline.models import Citation, MatterArtifacts


def _cite(cites: list[Citation]) -> str:
    return "; ".join(
        f"{c.file} p.{c.page}" + (f" ¶{c.para}" if c.para is not None else "") for c in cites
    )


def artifacts_to_docx(artifacts: MatterArtifacts) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    doc.add_heading(f"Hearing Brief — {artifacts.matter_id}", level=0)

    doc.add_heading("Chronology of Events", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["Date", "Event", "Source"]):
        table.rows[0].cells[i].text = h
    for ev in artifacts.chronology:
        row = table.add_row().cells
        row[0].text = ev.event_date.strftime("%d.%m.%Y") if ev.event_date else "Undated"
        row[1].text = ev.event + (" [verify: low-confidence OCR]" if ev.confidence == "low_ocr" else "")
        row[2].text = _cite(ev.cites)

    doc.add_heading("Chronology of Proceedings", level=1)
    for order in artifacts.proceedings:
        when = order.order_date.strftime("%d.%m.%Y") if order.order_date else "Undated"
        line = f"{when} — {order.direction}"
        if order.next_date:
            line += f" (next date: {order.next_date.strftime('%d.%m.%Y')})"
        doc.add_paragraph(f"{line}  [{_cite(order.cites)}]", style="List Bullet")

    doc.add_heading("Rival Contentions", level=1)
    for cont in artifacts.contentions:
        doc.add_paragraph(cont.issue, style="Heading 3")
        if cont.petitioner:
            doc.add_paragraph(
                f"Petitioner/Plaintiff: {cont.petitioner.position}  [{_cite(cont.petitioner.cites)}]"
            )
        if cont.respondent:
            doc.add_paragraph(
                f"Respondent/Defendant: {cont.respondent.position}  [{_cite(cont.respondent.cites)}]"
            )

    doc.add_heading("Issues for Determination", level=1)
    for issue in artifacts.issues:
        origin = "framed by court" if issue.origin == "framed_by_court" else "inferred"
        suffix = f"  [{_cite(issue.cites)}]" if issue.cites else ""
        doc.add_paragraph(f"{issue.text} ({origin}){suffix}", style="List Number")

    doc.add_heading("Document Index", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(["Exhibit", "Title", "Type", "Date", "OCR"]):
        table.rows[0].cells[i].text = h
    for entry in artifacts.doc_index:
        row = table.add_row().cells
        row[0].text = entry.exhibit_no or "—"
        row[1].text = entry.title
        row[2].text = entry.doc_type.value
        row[3].text = entry.doc_date.strftime("%d.%m.%Y") if entry.doc_date else "—"
        row[4].text = entry.ocr_quality

    if artifacts.conflicts:
        doc.add_heading("Conflicts on the Record", level=1)
        for conflict in artifacts.conflicts:
            doc.add_paragraph(conflict.fact, style="Heading 3")
            for pos in conflict.positions:
                doc.add_paragraph(f"{pos.position}  [{_cite(pos.cites)}]", style="List Bullet")

    if artifacts.not_found:
        doc.add_heading("Not Found in the Record", level=1)
        for item in artifacts.not_found:
            doc.add_paragraph(item, style="List Bullet")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
