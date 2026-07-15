"""Export a draft to court-format .docx."""

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from pipeline.drafting.models import DraftDocument, DraftParagraph, ListOfDatesEntry
from pipeline.models import Citation


def _cite_note(cites: list[Citation]) -> str:
    return "; ".join(
        f"{c.file} p.{c.page}" + (f" ¶{c.para}" if c.para is not None else "") for c in cites
    )


def _heading(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.underline = True


def _prose(
    doc,
    paragraphs: list[DraftParagraph],
    include_citation_notes: bool,
    numbered: bool = True,
) -> None:
    """Body prose: factual paragraphs numbered, grounds lettered (A., B., ...),
    headings bold, boilerplate plain. The synopsis passes numbered=False —
    convention there is unnumbered narrative."""
    n = 0
    g = 0
    for para in paragraphs:
        if para.kind == "heading":
            _heading(doc, para.text)
            continue
        if para.kind == "ground":
            g += 1
            p = doc.add_paragraph(f"{chr(64 + ((g - 1) % 26) + 1)}. {para.text}")
        elif para.kind == "factual" and numbered:
            n += 1
            p = doc.add_paragraph(f"{n}. {para.text}")
        else:
            p = doc.add_paragraph(para.text)
        p.paragraph_format.space_after = Pt(8)
        if include_citation_notes and para.cites:
            note = doc.add_paragraph(f"[record: {_cite_note(para.cites)}]")
            note.runs[0].font.size = Pt(8)
            note.runs[0].italic = True


def _list_of_dates(doc, entries: list[ListOfDatesEntry], include_citation_notes: bool) -> None:
    """The two-column Date | Event table. An undated event says so explicitly
    — a blank cell would read as an omission, not a bucket."""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    head = table.rows[0].cells
    head[0].paragraphs[0].add_run("DATE").bold = True
    head[1].paragraphs[0].add_run("EVENT").bold = True
    for e in entries:
        row = table.add_row().cells
        row[0].text = e.event_date.strftime("%d.%m.%Y") if e.event_date else "Undated"
        row[1].text = e.event
        if e.confidence == "low_ocr":
            note = row[1].add_paragraph("[low-confidence OCR — verify against the scan]")
            note.runs[0].font.size = Pt(8)
            note.runs[0].italic = True
        if include_citation_notes and e.cites:
            note = row[1].add_paragraph(f"[record: {_cite_note(e.cites)}]")
            note.runs[0].font.size = Pt(8)
            note.runs[0].italic = True


def draft_to_docx(draft: DraftDocument, include_citation_notes: bool = True) -> bytes:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    if draft.court_header:
        for line in draft.court_header.splitlines():
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(draft.title)
    run.bold = True
    run.underline = True

    # Paperbook front matter first, in filing order: synopsis, list of dates.
    if draft.synopsis:
        _heading(doc, "SYNOPSIS")
        _prose(doc, draft.synopsis, include_citation_notes, numbered=False)
    if draft.list_of_dates:
        _heading(doc, "LIST OF DATES & EVENTS")
        _list_of_dates(doc, draft.list_of_dates, include_citation_notes)
        if draft.paragraphs or draft.prayer:
            doc.add_page_break()  # the petition body starts on its own page

    _prose(doc, draft.paragraphs, include_citation_notes)

    if draft.prayer:
        heading = doc.add_paragraph()
        heading.add_run("PRAYER").bold = True
        for i, relief in enumerate(draft.prayer, 1):
            doc.add_paragraph(f"({chr(96 + i)}) {relief}")

    if draft.missing_info:
        doc.add_page_break()
        warn = doc.add_paragraph()
        warn.add_run(
            "DRAFTING NOTES — information not on the record (verify and fill):"
        ).bold = True
        for item in draft.missing_info:
            doc.add_paragraph(item, style="List Bullet")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
